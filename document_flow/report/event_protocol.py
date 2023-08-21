from odoo import _, models
from docx.shared import Pt, Mm
from docx.enum.text import WD_BREAK
from htmldocx import HtmlToDocx

import pytz


class EventProtocol(models.AbstractModel):
    _name = 'report.document_flow.event_protocol'
    _description = 'document_flow.event_protocol'
    _inherit = "report.report_docx.abstract"

    def generate_docx_report(self, doc, data, objs):
        event = objs[0]
        if event:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_after = 0
            paragraph = doc.add_paragraph(_('Protocol "%s" from %s')
                                          % (event.name, pytz.utc.localize(event.date_start).astimezone(
                                            pytz.timezone(self.env.user.tz) or pytz.utc).strftime('%d.%m.%Y')))
            run = paragraph.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            paragraph = doc.add_paragraph(_('Were present:'))
            paragraph.paragraph_format.space_before = Mm(5)
            run = paragraph.runs[0]
            run.font.italic = True
            run.font.italic = True
            for member in event.member_ids:
                doc.add_paragraph(member.name + ' - %s' % member.partner_id.function if member.partner_id else '')
            html_parser = HtmlToDocx()
            if event.question_ids:
                paragraph = doc.add_paragraph(_('Agenda of the event:'))
                paragraph.paragraph_format.space_before = Mm(5)
                run = paragraph.runs[0]
                run.font.italic = True
                run.font.underline = True
                counter = 1
                for question in event.question_ids:
                    cnt_prg = len(doc.paragraphs)
                    html_parser.add_html_to_document(question.name, doc)
                    doc.paragraphs[cnt_prg].runs[0].add_text('%s. ' % counter)
                    doc.paragraphs[cnt_prg].paragraph_format.space_before = Mm(5)
                    counter += 1
            paragraph = doc.add_paragraph(_('Decisions:'))
            paragraph.paragraph_format.space_before = Mm(5)
            run = paragraph.runs[0]
            run.font.italic = True
            run.font.underline = True
            for decision in event.decision_ids:
                cnt_prg = len(doc.paragraphs)
                html_parser.add_html_to_document(decision.name, doc)
                doc.paragraphs[cnt_prg].runs[0].add_text('%s. ' % decision.num)
                doc.paragraphs[cnt_prg].paragraph_format.space_before = Mm(5)
                if decision.responsible_id:
                    doc.add_paragraph(_('Responsible: %s') % decision.responsible_id.name)
                if decision.executor_ids:
                    doc.add_paragraph(_("Executors: %s", ', '.join(decision.executor_ids.mapped('name'))))
                if decision.date_deadline:
                    if decision.deadline_type == 'to_date':
                        doc.add_paragraph(_('Due date: %s') % decision.date_deadline.strftime('%d.%m.%Y'))
                    else:
                        doc.add_paragraph(_('Due date: within %s days after execution paragraph %s') % (
                            decision.number_days, decision.after_decision_id.num))
            if event.annex_ids:
                for annex in event.annex_ids:
                    doc.paragraphs[len(doc.paragraphs) - 1].runs[
                        len(doc.paragraphs[len(doc.paragraphs) - 1].runs) - 1].add_break(WD_BREAK.PAGE)
                    paragraph = doc.add_paragraph(_('Annex %s') % annex.num)
                    paragraph.paragraph_format.space_after = Mm(5)
                    run = paragraph.runs[0]
                    run.font.size = Pt(14)
                    run.font.bold = True
                    html_parser.add_html_to_document(annex.name, doc)
