from odoo import _, models
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from html2text import html2text


class ApprovalForm(models.AbstractModel):
    _name = 'report.contract.contract_approval_form'
    _description = 'contract.contract_approval_form'
    _inherit = "report.report_docx.abstract"

    def generate_docx_report(self, doc, data, objs):
        contract = objs[0]
        if contract:
            processes = contract.process_id.child_ids.filtered(lambda pr: pr.type == 'agreement')
            if any(processes):
                style = doc.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                style.paragraph_format.space_after = 0
                paragraph = doc.add_paragraph(_('Approval form'))
                paragraph.alignment = 1
                paragraph.paragraph_format.space_after = Mm(5)
                run = paragraph.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                paragraph = doc.add_paragraph(_('Contract "%s" from %s')
                                              % (contract.name, contract.date.strftime('%d.%m.%Y')))
                paragraph.paragraph_format.space_after = Mm(1)

                table = doc.add_table(rows=1, cols=5, style='TableGrid')
                head_cells = table.rows[0].cells
                for i, item in enumerate([_('Full name'), _('Actual executor'), _('Result'), _('Date'), _('Comment')]):
                    p = head_cells[i].paragraphs[0]
                    p.add_run(item).bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for task in processes.task_ids.filtered(
                        lambda t: not t.parent_id and t.stage_id.result_type != 'error'):
                    row = table.add_row().cells
                    row[0].text = task.role_executor_id.name if task.role_executor_id else ', '.join(
                        task.user_ids.mapped('name'))
                    row[1].text = task.actual_executor_id.name if task.actual_executor_id else ''
                    row[2].text = _('Agreed') if task.is_closed else _('Not agreed')
                    row[3].text = task.date_closed.strftime('%d.%m.%Y %H:%M:%S') if task.is_closed else ''
                    row[4].text = html2text(task.execution_result) if task.execution_result else ''
