from odoo import _, models
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import pytz


class ApprovalForm(models.AbstractModel):
    _name = 'report.document_flow.document.approval.form'
    _description = 'document_flow.document.approval.form'
    _inherit = 'report.report_docx.abstract'

    def generate_docx_report(self, doc, data, objs):
        document = objs[0]
        if document and document.workflow_process_id:
            activities = document.workflow_process_id.activity_ids.filtered(
                lambda act: not act.flow_start and not act.flow_stop and
                            act.activity_id.task_type_id.code == 'sys_df_agreement')
            if any(activities):
                style = doc.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                style.paragraph_format.space_after = 0
                paragraph = doc.add_paragraph(_('Approval form'))
                paragraph.alignment = 1
                paragraph.paragraph_format.space_after = Mm(5)
                run = paragraph.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                paragraph = doc.add_paragraph(_('Document "%s" from %s')
                                              % (document.name, document.date.strftime('%d.%m.%Y')))
                paragraph.paragraph_format.space_after = Mm(1)

                table = doc.add_table(rows=1, cols=5, style='TableGrid')
                head_cells = table.rows[0].cells
                for i, item in enumerate([_('Full name'), _('Actual executor'), _('Result'), _('Date'), _('Comment')]):
                    p = head_cells[i].paragraphs[0]
                    p.add_run(item).bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                tasks = self.env['task.task'].sudo().search([
                    ('activity_id', 'in', activities.ids)
                ])

                for activity in activities:
                    task = tasks.filtered(lambda t: t.activity_id.id == activity.id)[-1:]
                    row = table.add_row().cells
                    if task:
                        full_name = task.group_executors_id.name if task.group_executors_id else ', '.join(
                            task.user_ids.mapped('name'))
                    else:
                        full_name = activity.activity_id.group_executors_id.name \
                            if activity.activity_id.group_executors_id else ', '.join(
                                activity.activity_id.user_ids.mapped('name'))
                    row[0].text = full_name
                    row[1].text = task.actual_executor_id.name if task.actual_executor_id else ''
                    row[2].text = _('Agreed') if task and task.is_closed else _('Not agreed')
                    row[3].text = pytz.utc.localize(task.date_closed).astimezone(
                        pytz.timezone(self.env.user.tz) or pytz.utc).strftime(
                            '%d.%m.%Y') if task and task.is_closed else ''
                    row[4].text = task.execution_result_text if task and task.execution_result else ''
